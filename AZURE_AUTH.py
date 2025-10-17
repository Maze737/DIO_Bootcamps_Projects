import os
import re
import tempfile
import uuid
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
from docx import Document
from langdetect import detect, DetectorFactory

DetectorFactory.seed = 0

USER_AGENT = "Translate-AzureAuth/1.0"
REQUEST_TIMEOUT = 15
MAX_CHARS_PER_CHUNK = 3000

AZ_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
AZ_KEY = os.getenv("AZURE_OPENAI_KEY", "")
AZ_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT", "")

def _clean_text_from_html(html, paragraph_sep="\n\n"):
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    raw = soup.get_text(separator="\n")
    lines = [line.strip() for line in raw.splitlines()]
    non_empty = [re.sub(r'\s+', ' ', line) for line in lines if line]
    return paragraph_sep.join(non_empty)

def _validate_url(url):
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("URL must use http or https")
    return url

def extract_text_from_url(url):
    _validate_url(url)
    headers = {"User-Agent": USER_AGENT}
    r = requests.get(url, headers=headers, timeout=REQUEST_TIMEOUT)
    r.raise_for_status()
    return _clean_text_from_html(r.text)

def detect_language(text):
    try:
        return detect(text)
    except Exception:
        return "en"

def _azure_chat_call(prompt):
    if not (AZ_ENDPOINT and AZ_KEY and AZ_DEPLOYMENT):
        raise RuntimeError("Set AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_KEY and AZURE_OPENAI_DEPLOYMENT")
    url = f"{AZ_ENDPOINT}/openai/deployments/{AZ_DEPLOYMENT}/chat/completions?api-version=2024-06-01-preview"
    headers = {"api-key": AZ_KEY, "Content-Type": "application/json"}
    payload = {
        "messages": [
            {"role": "system", "content": "You are a precise technical translator. Preserve terminology and domain context."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 2000
    }
    r = requests.post(url, headers=headers, json=payload, timeout=120)
    r.raise_for_status()
    j = r.json()
    return j["choices"][0]["message"]["content"]

def _chunked_azure_translate(text, src, tgt, as_markdown=False):
    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    out = []
    buffer = []
    buf_len = 0
    for p in paragraphs:
        if buf_len + len(p) + 1 > MAX_CHARS_PER_CHUNK and buffer:
            chunk = "\n\n".join(buffer)
            prompt = f"Source language: {src}. Translate to {tgt}. Preserve technical terms. {'Respond in Markdown.' if as_markdown else 'Respond only with translated text.'}\n\n{chunk}"
            out.append(_azure_chat_call(prompt))
            buffer = [p]; buf_len = len(p)
        else:
            buffer.append(p); buf_len += len(p) + 1
    if buffer:
        chunk = "\n\n".join(buffer)
        prompt = f"Source language: {src}. Translate to {tgt}. Preserve technical terms. {'Respond in Markdown.' if as_markdown else 'Respond only with translated text.'}\n\n{chunk}"
        out.append(_azure_chat_call(prompt))
    return "\n\n".join(out)

def translate_url_to_markdown(url, target_lang="pt-BR"):
    text = extract_text_from_url(url)
    src = detect_language(text)
    return _chunked_azure_translate(text, src, target_lang, as_markdown=True)

def translate_text_to_docx(text, target_lang="pt-BR"):
    src = detect_language(text)
    translated = _chunked_azure_translate(text, src, target_lang, as_markdown=False)
    tmp = tempfile.NamedTemporaryFile(delete=False, prefix=f"translated_{uuid.uuid4().hex}_", suffix=".docx")
    doc = Document()
    for line in translated.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(tmp.name)
    tmp.close()
    return tmp.name

def translate_docx(path, target_lang="pt-BR"):
    doc = Document(path)
    full = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return translate_text_to_docx(full, target_lang)
