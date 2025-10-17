import os
import re
import tempfile
import uuid
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
from docx import Document
from langdetect import detect, DetectorFactory
from transformers import pipeline

DetectorFactory.seed = 0

USER_AGENT = "Translate-GeneralUse/1.0 (+https://example.com)"
REQUEST_TIMEOUT = 15
MAX_CHARS_PER_CHUNK = 3500  # safe chunk size for local models

def _clean_text_from_html(html, paragraph_sep="\n\n"):
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    raw = soup.get_text(separator="\n")
    lines = [line.strip() for line in raw.splitlines()]
    non_empty = [re.sub(r'\s+', ' ', line) for line in lines if line]
    return paragraph_sep.join(non_empty)

def _validate_url(url):
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("URL must use http or https")
    return url

def extract_text_from_url(url):
    _validate_url(url)
    headers = {"User-Agent": USER_AGENT}
    r = requests.get(url, headers=headers, timeout=REQUEST_TIMEOUT)
    r.raise_for_status()
    return _clean_text_from_html(r.text)

def detect_language(text):
    try:
        return detect(text)
    except Exception:
        return "en"

def _get_translator(src_lang, tgt_lang):
    # maps codes like 'pt' or 'en' to model id used by Helsinki-NLP
    src = src_lang.split('-')[0]
    tgt = tgt_lang.split('-')[0]
    model_id = f"Helsinki-NLP/opus-mt-{src}-{tgt}"
    try:
        return pipeline("translation", model=model_id, device=-1)
    except Exception:
        # fallback to en->tgt
        model_id = f"Helsinki-NLP/opus-mt-en-{tgt}"
        return pipeline("translation", model=model_id, device=-1)

def _chunked_translate(text, src, tgt):
    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    translator = _get_translator(src, tgt)
    out_parts = []
    buffer = []
    buf_len = 0
    for p in paragraphs:
        if buf_len + len(p) + 1 > MAX_CHARS_PER_CHUNK and buffer:
            chunk = "\n\n".join(buffer)
            res = translator(chunk, max_length=4000)
            out_parts.append(res[0]["translation_text"])
            buffer = [p]; buf_len = len(p)
        else:
            buffer.append(p); buf_len += len(p) + 1
    if buffer:
        chunk = "\n\n".join(buffer)
        res = translator(chunk, max_length=4000)
        out_parts.append(res[0]["translation_text"])
    return "\n\n".join(out_parts)

def translate_url_to_markdown(url, target_lang="pt"):
    text = extract_text_from_url(url)
    src = detect_language(text)
    translated = _chunked_translate(text, src, target_lang)
    # transformers return plain text; user asked URL->markdown string: assume translator preserves paragraphs
    return translated

def translate_text_to_docx(text, target_lang="pt"):
    src = detect_language(text)
    translated = _chunked_translate(text, src, target_lang)
    tmp = tempfile.NamedTemporaryFile(delete=False, prefix=f"translated_{uuid.uuid4().hex}_", suffix=".docx")
    doc = Document()
    for line in translated.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(tmp.name)
    tmp.close()
    return tmp.name

def translate_docx(path, target_lang="pt"):
    doc = Document(path)
    full = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return translate_text_to_docx(full, target_lang)
