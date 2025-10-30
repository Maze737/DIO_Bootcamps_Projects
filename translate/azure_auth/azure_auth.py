"""
Módulo de tradução autenticada via Azure OpenAI.

Oferece funções simples para traduzir textos técnicos vindos de URLs, de
arquivos .docx ou diretamente do terminal, preservando termos de domínio.
"""

from __future__ import annotations

import os
import tempfile
import uuid
from typing import Iterable, List
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from langdetect import DetectorFactory, detect

DetectorFactory.seed = 0  # resulta em detecção reprodutível

USER_AGENT = "Translate-AzureAuth/1.0"
REQUEST_TIMEOUT = 15
MAX_CHARS_PER_CHUNK = 2800
MAX_DOWNLOAD_BYTES = 2_500_000

AZ_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
AZ_KEY = os.getenv("AZURE_OPENAI_KEY", "")
AZ_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT", "")
AZ_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-06-01-preview")


def clean_text(texto: str) -> str:
    """Remove espaços extras e linhas vazias preservando legibilidade."""
    # Limpar texto
    linhas = (line.strip() for line in texto.splitlines())
    parts = (phrase.strip() for line in linhas for phrase in line.split(" "))
    texto_limpo = " ".join(part for part in parts if part)
    return texto_limpo


def _normalize_blocks(texto: str) -> str:
    blocos = []
    for bloco in texto.split("\n\n"):
        trecho = clean_text(bloco)
        if trecho:
            blocos.append(trecho)
    return "\n\n".join(blocos)


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    raw = soup.get_text(separator="\n")
    return _normalize_blocks(raw)


def _validate_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("URL precisa utilizar http ou https.")
    if not parsed.netloc:
        raise ValueError("URL inválida: domínio não encontrado.")
    return url


def extract_text_from_url(url: str) -> str:
    """Obtém e sanitiza o texto principal de uma página web."""
    _validate_url(url)
    headers = {"User-Agent": USER_AGENT}
    response = requests.get(url, headers=headers, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    if len(response.content) > MAX_DOWNLOAD_BYTES:
        raise ValueError("Conteúdo muito extenso para tradução automática.")
    html = response.content.decode(response.encoding or "utf-8", errors="ignore")
    return _html_to_text(html)


def detect_language(texto: str) -> str:
    texto_limpo = clean_text(texto)
    try:
        return detect(texto_limpo) if texto_limpo else "en"
    except Exception:
        return "en"


def _ensure_azure_settings() -> None:
    if not (AZ_ENDPOINT and AZ_KEY and AZ_DEPLOYMENT):
        raise RuntimeError(
            "Configure AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_KEY e AZURE_OPENAI_DEPLOYMENT."
        )


def _chunk_text(texto: str, limite: int) -> Iterable[str]:
    paragrafos = [p for p in texto.split("\n\n") if p.strip()]
    buffer: List[str] = []
    tamanho = 0
    for paragrafo in paragrafos:
        if tamanho + len(paragrafo) + 1 > limite and buffer:
            yield "\n\n".join(buffer)
            buffer = [paragrafo]
            tamanho = len(paragrafo)
        else:
            buffer.append(paragrafo)
            tamanho += len(paragrafo) + 1
    if buffer:
        yield "\n\n".join(buffer)


def _azure_chat_call(prompt: str) -> str:
    _ensure_azure_settings()
    url = (
        f"{AZ_ENDPOINT}/openai/deployments/{AZ_DEPLOYMENT}"
        f"/chat/completions?api-version={AZ_API_VERSION}"
    )
    headers = {
        "api-key": AZ_KEY,
        "Content-Type": "application/json",
        "User-Agent": USER_AGENT,
    }
    payload = {
        "messages": [
            {
                "role": "system",
                "content": (
                    "Você é um tradutor técnico preciso. Preserve termos de domínio, "
                    "abreviações, unidades e formatação relevante."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0,
        "max_tokens": 1800,
    }
    response = requests.post(url, headers=headers, json=payload, timeout=120)
    response.raise_for_status()
    data = response.json()
    try:
        return data["choices"][0]["message"]["content"]
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Formato inesperado da resposta Azure: {data}") from exc


def _translate_chunks(chunks: Iterable[str], origem: str, destino: str, markdown: bool) -> str:
    partes = []
    for trecho in chunks:
        prompt = (
            f"Idioma de origem: {origem}. Traduza para {destino}. "
            "Mantenha o contexto técnico, números e nomenclatura. "
            f"{'Responda em markdown padronizado.' if markdown else 'Responda apenas com o texto traduzido.'}\n\n"
            f"{trecho}"
        )
        partes.append(_azure_chat_call(prompt))
    return "\n\n".join(partes)


def translate_text(texto: str, target_lang: str = "pt-BR") -> str:
    """Traduz texto bruto (ex.: digitado no terminal) e retorna string limpa."""
    normalizado = _normalize_blocks(texto)
    origem = detect_language(normalizado)
    return _translate_chunks(
        _chunk_text(normalizado, MAX_CHARS_PER_CHUNK),
        origem,
        target_lang,
        markdown=False,
    )


def translate_url_to_markdown(url: str, target_lang: str = "pt-BR") -> str:
    """Traduz o conteúdo de uma página e devolve markdown pronto para uso."""
    texto = extract_text_from_url(url)
    origem = detect_language(texto)
    return _translate_chunks(
        _chunk_text(texto, MAX_CHARS_PER_CHUNK),
        origem,
        target_lang,
        markdown=True,
    )


def translate_text_to_docx(texto: str, target_lang: str = "pt-BR") -> str:
    """Traduz texto livre e salva uma cópia em DOCX; retorna o caminho final."""
    traducao = translate_text(texto, target_lang)
    destino = tempfile.NamedTemporaryFile(
        delete=False, prefix=f"translated_{uuid.uuid4().hex}_", suffix=".docx"
    )
    destino.close()
    doc = Document()
    for linha in traducao.splitlines():
        linha_limpa = linha.strip()
        if linha_limpa:
            doc.add_paragraph(linha_limpa)
    doc.save(destino.name)
    return destino.name


def translate_docx(path: str, target_lang: str = "pt-BR") -> str:
    """Traduz um arquivo DOCX existente e devolve o caminho do novo arquivo."""
    doc = Document(path)
    texto = "\n\n".join(
        paragrafo.text for paragrafo in doc.paragraphs if paragrafo.text.strip()
    )
    return translate_text_to_docx(texto, target_lang)


__all__ = [
    "clean_text",
    "detect_language",
    "extract_text_from_url",
    "translate_text",
    "translate_docx",
    "translate_url_to_markdown",
    "translate_text_to_docx",
]
