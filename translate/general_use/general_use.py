"""
Módulo de uso geral com tradução offline usando Hugging Face.

Entrega funções simples para traduzir textos técnicos sem autenticação,
aceitando URLs, arquivos .docx ou entrada direta no terminal.
"""

from __future__ import annotations

import functools
import tempfile
import uuid
from typing import Iterable, List
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from langdetect import DetectorFactory, detect
from transformers import pipeline

DetectorFactory.seed = 0

USER_AGENT = "Translate-GeneralUse/1.0"
REQUEST_TIMEOUT = 15
MAX_CHARS_PER_CHUNK = 1800
MAX_DOWNLOAD_BYTES = 1_500_000


def clean_text(texto: str) -> str:
    """Remove espaços extras e linhas vazias preservando legibilidade."""
    # Limpar texto
    linhas = (line.strip() for line in texto.splitlines())
    parts = (phrase.strip() for line in linhas for phrase in line.split(" "))
    texto_limpo = " ".join(part for part in parts if part)
    return texto_limpo


def _normalize_blocks(texto: str) -> str:
    blocos = []
    for bloco in texto.split("\n\n"):
        trecho = clean_text(bloco)
        if trecho:
            blocos.append(trecho)
    return "\n\n".join(blocos)


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    raw = soup.get_text(separator="\n")
    return _normalize_blocks(raw)


def _validate_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("URL precisa utilizar http ou https.")
    if not parsed.netloc:
        raise ValueError("URL inválida: domínio não encontrado.")
    return url


def extract_text_from_url(url: str) -> str:
    _validate_url(url)
    headers = {"User-Agent": USER_AGENT}
    response = requests.get(url, headers=headers, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    if len(response.content) > MAX_DOWNLOAD_BYTES:
        raise ValueError("Conteúdo muito extenso para tradução automática.")
    html = response.content.decode(response.encoding or "utf-8", errors="ignore")
    return _html_to_text(html)


def detect_language(texto: str) -> str:
    texto_limpo = clean_text(texto)
    try:
        return detect(texto_limpo) if texto_limpo else "en"
    except Exception:
        return "en"


def _chunk_text(texto: str, limite: int) -> Iterable[str]:
    paragrafos = [p for p in texto.split("\n\n") if p.strip()]
    buffer: List[str] = []
    tamanho = 0
    for paragrafo in paragrafos:
        if tamanho + len(paragrafo) + 1 > limite and buffer:
            yield "\n\n".join(buffer)
            buffer = [paragrafo]
            tamanho = len(paragrafo)
        else:
            buffer.append(paragrafo)
            tamanho += len(paragrafo) + 1
    if buffer:
        yield "\n\n".join(buffer)


def _base_lang(codigo: str) -> str:
    return codigo.split('-')[0] if codigo else ''


@functools.lru_cache(maxsize=8)
def _load_pipeline(model_id: str):
    return pipeline("translation", model=model_id, device=-1)


def _get_translator(origem: str, destino: str):
    src = _base_lang(origem) or "en"
    tgt = _base_lang(destino) or "pt"
    candidatos = [
        f"Helsinki-NLP/opus-mt-{src}-{tgt}",
        f"Helsinki-NLP/opus-mt-en-{tgt}",
        f"Helsinki-NLP/opus-mt-mul-{tgt}",
        "Helsinki-NLP/opus-mt-mul-en",
    ]
    erros = []
    for modelo in candidatos:
        try:
            return _load_pipeline(modelo)
        except Exception as exc:  # pragma: no cover - depende do ambiente
            erros.append(f"{modelo}: {exc}")
    raise RuntimeError(
        "Nenhum modelo Hugging Face pôde ser carregado. " + "; ".join(erros)
    )


def _translate_chunks(texto: str, origem: str, destino: str) -> str:
    translator = _get_translator(origem, destino)
    partes = []
    for trecho in _chunk_text(texto, MAX_CHARS_PER_CHUNK):
        resultado = translator(trecho, max_length=4000)
        partes.append(resultado[0]["translation_text"])
    return "\n\n".join(partes)


def translate_text(texto: str, target_lang: str = "pt") -> str:
    normalizado = _normalize_blocks(texto)
    origem = detect_language(normalizado)
    return _translate_chunks(normalizado, origem, target_lang)


def translate_url_to_markdown(url: str, target_lang: str = "pt") -> str:
    texto = extract_text_from_url(url)
    origem = detect_language(texto)
    return _translate_chunks(texto, origem, target_lang)


def translate_text_to_docx(texto: str, target_lang: str = "pt") -> str:
    traducao = translate_text(texto, target_lang)
    destino = tempfile.NamedTemporaryFile(
        delete=False, prefix=f"translated_{uuid.uuid4().hex}_", suffix=".docx"
    )
    destino.close()
    doc = Document()
    for linha in traducao.splitlines():
        linha_limpa = linha.strip()
        if linha_limpa:
            doc.add_paragraph(linha_limpa)
    doc.save(destino.name)
    return destino.name


def translate_docx(path: str, target_lang: str = "pt") -> str:
    doc = Document(path)
    texto = "\n\n".join(
        paragrafo.text for paragrafo in doc.paragraphs if paragrafo.text.strip()
    )
    return translate_text_to_docx(texto, target_lang)


__all__ = [
    "clean_text",
    "detect_language",
    "extract_text_from_url",
    "translate_text",
    "translate_docx",
    "translate_url_to_markdown",
    "translate_text_to_docx",
]
